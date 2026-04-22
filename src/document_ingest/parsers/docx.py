# src/document_ingest/parsers/docx.py
"""DOCX 파서 — python-docx 기반 (M_06 스펙 §5.2)."""

from __future__ import annotations

import logging
from pathlib import Path

from document_ingest.errors import ParseError
from document_ingest.segments import _Segment

logger = logging.getLogger(__name__)


def _parse_docx(path: str) -> list[_Segment]:
    """DOCX 파일에서 단락·표 텍스트 세그먼트를 추출한다.

    - section: 가장 최근 Heading 스타일 단락의 text.
    - page: DOCX 특성상 결정 불가 → 항상 None.
    - 표: 각 행의 셀 텍스트를 " | "로 연결해 1 세그먼트.

    Args:
        path: 절대 또는 상대 경로 문자열.

    Returns:
        list[_Segment]

    Raises:
        ParseError: python-docx 로드 실패.
    """
    # XXE 방어: defusedxml.defuse_stdlib() (A-3 스펙 §11.4)
    try:
        import defusedxml

        defusedxml.defuse_stdlib()
    except ImportError:
        logger.warning("defusedxml 미설치 — XXE 방어 비활성화")

    try:
        import docx
    except ImportError as exc:
        raise ParseError(f"python-docx 패키지가 설치되지 않았습니다: {exc}") from exc

    try:
        doc = docx.Document(path)
    except Exception as exc:
        raise ParseError(f"DOCX 열기 실패: {Path(path).name}: {exc}") from exc

    segments: list[_Segment] = []
    current_heading: str | None = None

    for para in doc.paragraphs:
        style_name: str = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            heading_text = para.text.strip()
            if heading_text:
                current_heading = heading_text

        text = para.text.strip()
        if text:
            segments.append(
                _Segment(
                    text=para.text,
                    page=None,
                    section=current_heading,
                    bbox=None,
                )
            )

    # 표: para 순회에서 잡히지 않으므로 별도 순회
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                segments.append(
                    _Segment(
                        text=row_text,
                        page=None,
                        section=current_heading,
                        bbox=None,
                    )
                )

    return segments
